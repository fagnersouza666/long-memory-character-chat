import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import openai
from dotenv import load_dotenv
import tempfile

# Load environment variables
load_dotenv()

# Configure page
st.set_page_config(
    page_title="Gerador de Documentos",
    page_icon="📄",
    layout="centered"
)

# Custom CSS for modern look
st.markdown("""
<style>
    .stApp {
        background-color: #f5f7fa;
    }
    .main-header {
        text-align: center;
        color: #2c3e50;
        margin-bottom: 2rem;
    }
    .stButton>button {
        background-color: #3498db;
        color: white;
        border-radius: 8px;
        border: none;
        padding: 0.5rem 1rem;
        font-size: 1rem;
        transition: all 0.3s;
    }
    .stButton>button:hover {
        background-color: #2980b9;
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    .stTextArea textarea {
        border-radius: 8px;
        border: 1px solid #ddd;
    }
    .stFileUploader {
        border: 2px dashed #3498db;
        border-radius: 8px;
        padding: 1rem;
        text-align: center;
    }
    .template-card {
        border: 1px solid #ddd;
        border-radius: 8px;
        padding: 1rem;
        margin-bottom: 1rem;
        background-color: white;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        cursor: pointer;
        transition: all 0.3s;
    }
    .template-card:hover {
        transform: translateY(-3px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.15);
    }
    .template-card.selected {
        border-color: #3498db;
        background-color: #e1f0fa;
    }
    .output-preview {
        background-color: white;
        border-radius: 8px;
        padding: 1.5rem;
        margin-top: 1.5rem;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if "selected_template" not in st.session_state:
    st.session_state.selected_template = None
if "generated_content" not in st.session_state:
    st.session_state.generated_content = ""
if "file_name" not in st.session_state:
    st.session_state.file_name = ""

# Main header
st.markdown("<h1 class='main-header'>📄 Gerador de Documentos Simples</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; color: #7f8c8d;'>Crie documentos profissionais com facilidade</p>", unsafe_allow_html=True)

# Function to extract text from uploaded files
def extract_text_from_file(uploaded_file):
    """Extract text from uploaded file (PDF, DOCX, TXT)"""
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    if file_extension == "txt":
        return uploaded_file.getvalue().decode("utf-8")
    elif file_extension == "docx":
        from docx import Document
        # Use a temporary file to handle the uploaded file
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        
        try:
            doc = Document(tmp_file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        finally:
            os.unlink(tmp_file_path)
    elif file_extension == "pdf":
        import PyPDF2
        # Use a temporary file to handle the uploaded file
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        
        try:
            with open(tmp_file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text
        finally:
            os.unlink(tmp_file_path)
    else:
        st.error("Formato de arquivo não suportado. Por favor, use TXT, DOCX ou PDF.")
        return None

# Function to call GPT-5-mini
def process_with_gpt(content, template):
    """Process content with GPT-5-mini according to selected template"""
    try:
        # Get API key from environment
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            st.warning("Chave API não encontrada. O processamento será simulado.")
            # Simulate processing for demo purposes
            return f"[Conteúdo processado com o modelo {template}]\n\n{content}"
            
        client = openai.OpenAI(api_key=api_key)
        
        # Create prompt based on template
        prompt = f"""
        Formate o seguinte conteúdo de acordo com o modelo '{template}':
        
        {content}
        
        Instruções:
        - Mantenha o conteúdo essencial
        - Aplique a formatação apropriada para o modelo selecionado
        - Seja conciso e profissional
        - Não adicione informações que não estejam no conteúdo original
        """
        
        # Call GPT-5-mini with fixed parameters for consistency
        response = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {"role": "system", "content": "Você é um assistente especializado em formatação de documentos. Formate o conteúdo de acordo com o modelo solicitado."},
                {"role": "user", "content": prompt}
            ],
            max_completion_tokens=2000
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        st.warning(f"Processamento simulado devido a: {str(e)}")
        # Simulate processing for demo purposes
        return f"[Conteúdo processado com o modelo {template}]\n\n{content}"

# Function to create DOCX document
def create_docx(content, template):
    """Create DOCX document with formatted content"""
    doc = Document()
    
    # Add title based on template
    title = doc.add_heading(template, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content
    for paragraph in content.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

# Function to create PDF document
def create_pdf(content, template):
    """Create PDF document with formatted content"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            alignment=TA_CENTER,
            fontSize=24,
            spaceAfter=30,
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=12,
        )
        
        # Build document content
        flowables = []
        
        # Add title
        title = Paragraph(template, title_style)
        flowables.append(title)
        flowables.append(Spacer(1, 0.5*inch))
        
        # Add content
        for paragraph in content.split('\n'):
            if paragraph.strip():
                para = Paragraph(paragraph, normal_style)
                flowables.append(para)
                flowables.append(Spacer(1, 0.2*inch))
        
        # Build PDF
        doc.build(flowables)
        buffer.seek(0)
        
        return buffer
    except Exception as e:
        st.error(f"Erro ao criar PDF: {str(e)}")
        return None

# Input section
st.subheader("📝 Entrada de Conteúdo")
input_method = st.radio("Escolha o método de entrada:", ["Texto Direto", "Upload de Arquivo"])

content = ""
if input_method == "Texto Direto":
    content = st.text_area("Digite seu conteúdo aqui:", height=200, max_chars=5000,
                          placeholder="Cole ou digite seu conteúdo aqui...")
else:
    uploaded_file = st.file_uploader("Faça upload de um arquivo (TXT, DOCX, PDF):", 
                                     type=["txt", "docx", "pdf"])
    if uploaded_file:
        content = extract_text_from_file(uploaded_file)
        if content:
            st.success("Arquivo carregado com sucesso!")

# Template selection
st.subheader("📋 Escolha um Modelo")

# Define templates
templates = {
    "Relatório Empresarial": {
        "description": "Formato profissional para relatórios corporativos",
        "icon": "📊"
    },
    "Documento Técnico": {
        "description": "Estrutura técnica com seções e formatação específica",
        "icon": "⚙️"
    },
    "Carta/Memorando": {
        "description": "Formato formal para comunicações oficiais",
        "icon": "✉️"
    },
    "Resumo Executivo": {
        "description": "Versão concisa com pontos principais",
        "icon": "📋"
    },
    "Documento Criativo": {
        "description": "Formato flexível para conteúdo criativo",
        "icon": "✏️"
    }
}

# Display templates as cards
cols = st.columns(2)
template_names = list(templates.keys())

for i, (name, info) in enumerate(templates.items()):
    col = cols[i % 2]
    with col:
        is_selected = st.session_state.selected_template == name
        button_type = "primary" if is_selected else "secondary"
        
        if st.button(f"{info['icon']} {name}", key=f"template_{i}", 
                     type=button_type, use_container_width=True):
            st.session_state.selected_template = name
            st.rerun()
        
        st.caption(info["description"])

# Generate button
st.subheader("🚀 Gerar Documento")
if st.button("Gerar Documento", type="primary", use_container_width=True, 
             disabled=not content or not st.session_state.selected_template):
    
    if content and st.session_state.selected_template:
        with st.spinner("Processando com GPT-5-mini..."):
            # Process content with GPT-5-mini
            processed_content = process_with_gpt(content, st.session_state.selected_template)
            
            if processed_content:
                st.session_state.generated_content = processed_content
                st.session_state.file_name = f"documento_{st.session_state.selected_template.lower().replace(' ', '_')}"
                st.success("Documento gerado com sucesso!")
            else:
                st.error("Falha ao gerar o documento. Por favor, tente novamente.")

# Output section
if st.session_state.generated_content:
    st.subheader("💾 Resultado")
    
    # Preview
    st.markdown("<div class='output-preview'>", unsafe_allow_html=True)
    st.markdown("### Pré-visualização")
    st.text_area("Conteúdo formatado:", value=st.session_state.generated_content, height=200, 
                 key="preview_content")
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Download options
    st.markdown("### Baixar Documento")
    col1, col2 = st.columns(2)
    
    with col1:
        # Download as DOCX
        docx_buffer = create_docx(st.session_state.generated_content, st.session_state.selected_template)
        st.download_button(
            label="📥 Baixar como DOCX",
            data=docx_buffer,
            file_name=f"{st.session_state.file_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    with col2:
        # Download as PDF
        pdf_buffer = create_pdf(st.session_state.generated_content, st.session_state.selected_template)
        if pdf_buffer:
            st.download_button(
                label="📥 Baixar como PDF",
                data=pdf_buffer,
                file_name=f"{st.session_state.file_name}.pdf",
                mime="application/pdf"
            )
        else:
            st.warning("Não foi possível gerar o PDF. Verifique os logs para mais detalhes.")

# Footer
st.markdown("---")
st.caption("Todos os documentos são processados com GPT-5-mini exclusivamente. O conteúdo é gerado apenas para fins de demonstração.")